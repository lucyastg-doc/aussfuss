from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Dict, Any, Tuple
from datetime import date
from docx import Document


def _cat_size_mm(mm: float) -> str:
    if mm <= 2:
        return "<=2"
    if mm <= 4:
        return "2-4"
    if mm <= 6:
        return "4-6"
    return ">6"


def _cat_min_thickness(min_um: int) -> str:
    if min_um >= 400:
        return ">=400"
    if 300 <= min_um <= 399:
        return "300-399"
    if 200 <= min_um <= 299:
        return "200-299"
    return "<200"


def _cat_mean_thickness(mean_um: int) -> str:
    if mean_um >= 600:
        return ">=600"
    if 520 <= mean_um <= 599:
        return "520-599"
    if 450 <= mean_um <= 519:
        return "450-519"
    return "<450"


@dataclass(frozen=True)
class ScoreResult:
    score: int
    breakdown: Dict[str, int]
    critical: bool


def compute_fuss(ctx: Dict[str, Any]) -> ScoreResult:
    bd: Dict[str, int] = {}

    # Общие клинические признаки (в начале — одинаково в обеих шкалах)
    bd["Болевой синдром"] = {0: 0, 2: 2, 4: 4}[ctx.get("pain", 0)]
    bd["Перикорнеальная инъекция"] = {0: 0, 1: 1, 2: 2, 3: 3}[ctx.get("injection", 0)]
    bd["Отделяемое"] = {0: 0, 1: 2}[ctx.get("discharge", 0)]
    bd["Сателлитные инфильтраты/«перистые» края"] = {0: 0, 1: 2}[ctx.get("satellites", 0)]

    # Размер/локализация/глубина
    bd["Размер дефекта"] = {"<=2": 0, "2-4": 1, "4-6": 2, ">6": 3}[ctx["size_cat"]]
    bd["Клиническая форма (грибковая)"] = {0: 0, 1: 1, 2: 2}[ctx.get("fungal_form", 0)]
    bd["Локализация"] = {"peripheral": 0, "paracentral": 1, "central": 2}[ctx["localization"]]
    bd["Глубина"] = {"superficial": 0, "mid": 2, "deep": 4, "descemetocele": 6}[ctx["depth_cat"]]

    # Воспаление/ПК/гипопион/отёк
    bd["Признаки десцеметита"] = {0: 0, 1: 2}[ctx.get("descemetitis", 0)]
    bd["Гипопион"] = {"none": 0, "lt1": 1, "1to2": 2, "gt2": 3}[ctx.get("hypopyon", "none")]
    bd["Тотальное бельмо"] = {0: 0, 1: 2}[ctx.get("total_leucoma", 0)]

    # ЕДИНЫЙ вопрос по передней камере для обеих шкал
    # Для FUSS учитывается «не просматривается» отдельно (4 балла)
    ac = ctx.get("ac", "0")
    bd["Передняя камера"] = {"0": 0, "1-20": 1, ">20": 2, "not_visible": 4}[ac]

    bd["Отёк роговицы"] = {0: 0, 1: 1, 2: 2}[ctx.get("edema", 0)]
    bd["ВГД"] = {"normal": 0, "high": 1, "low": 1}[ctx.get("iog", "normal")]

    # ОКТ — только пахиметрия
    bd["ОКТ: локальные зоны истончения"] = {0: 0, 1: 2}[ctx.get("pachy_uneven", 0)]
    min_cat = _cat_min_thickness(int(ctx.get("min_thickness_um", 400)))
    bd["ОКТ: минимальная толщина"] = {">=400": 0, "300-399": 2, "200-299": 4, "<200": 6}[min_cat]
    mean_cat = _cat_mean_thickness(int(ctx.get("mean_thickness_um", 600)))
    bd["ОКТ: средняя толщина"] = {">=600": 0, "520-599": 1, "450-519": 2, "<450": 3}[mean_cat]
    bd["ОКТ: прогрессирование истончения 48–72 ч"] = {0: 0, 1: 2}[ctx.get("thinning_progress_72h", 0)]

    # Лимб (как в AUSS — 0/2)
    bd["Вовлечение лимба"] = {0: 0, 1: 2}[ctx.get("limbal", 0)]

    # Конфокальная микроскопия — только конфокальная
    bd["Конфокальная: гифы"] = {0: 0, 1: 3, 2: 6, 3: 9}[ctx.get("hyphae", 0)]
    bd["Конфокальная: глубина гиф/спор"] = {0: 0, 1: 2, 2: 4, 3: 6}[ctx.get("hyphae_depth", 0)]

    # Скорость/прогноз (по клинике) — как было
    bd["Скорость прогрессирования"] = {0: 0, 1: 2, 2: 4}[ctx.get("progress_speed", 0)]
    bd["Прогноз интенсивности помутнения"] = {0: 0, 1: 1, 2: 2}[ctx.get("opacity", 0)]

    score = int(sum(bd.values()))
    critical = (min_cat == "<200") or (ctx.get("depth_cat") == "descemetocele")
    return ScoreResult(score=score, breakdown=bd, critical=critical)


def compute_auss(ctx: Dict[str, Any]) -> ScoreResult:
    bd: Dict[str, int] = {}

    bd["Болевой синдром"] = {0: 0, 2: 2, 4: 4}[ctx.get("pain", 0)]
    bd["Перикорнеальная инъекция"] = {0: 0, 1: 1, 2: 2, 3: 3}[ctx.get("injection", 0)]
    bd["Отделяемое"] = {0: 0, 1: 2}[ctx.get("discharge", 0)]
    bd["Сателлитные инфильтраты/«перистые» края"] = {0: 0, 1: 2}[ctx.get("satellites", 0)]

    # Специфическая клиника AUSS
    bd["Клиническая форма (AUSS)"] = {0: 0, 1: 1, 2: 2, 3: 3, 4: 4}[ctx.get("amoeba_form", 0)]
    bd["Эпителиальный дефект/псевдодендрит"] = {0: 0, 1: 1}[ctx.get("pseudo_dendrite", 0)]
    bd["Кольцевидный инфильтрат"] = {0: 0, 1: 3, 2: 6}[ctx.get("ring", 0)]
    bd["Радиальный кератоневрит (клиника)"] = {0: 0, 1: 2}[ctx.get("rk_clin", 0)]

    bd["Размер дефекта"] = {"<=2": 0, "2-4": 1, "4-6": 2, ">6": 3}[ctx["size_cat"]]
    bd["Локализация"] = {"peripheral": 0, "paracentral": 1, "central": 2}[ctx["localization"]]
    bd["Признаки десцеметита"] = {0: 0, 1: 2}[ctx.get("descemetitis", 0)]
    bd["Глубина"] = {"superficial": 0, "mid": 2, "deep": 4, "descemetocele": 6}[ctx["depth_cat"]]

    bd["Гипопион"] = {"none": 0, "lt1": 1, "1to2": 2, "gt2": 3}[ctx.get("hypopyon", "none")]
    bd["Тотальное бельмо"] = {0: 0, 1: 2}[ctx.get("total_leucoma", 0)]

    # ЕДИНЫЙ вопрос по передней камере
    # Для AUSS «не просматривается» приравниваем к >20 клеток (2 балла), чтобы не ломать единый интерфейс.
    ac = ctx.get("ac", "0")
    ac_a = {"0": 0, "1-20": 1, ">20": 2, "not_visible": 2}[ac]
    bd["Передняя камера"] = ac_a

    bd["Отёк роговицы"] = {0: 0, 1: 1, 2: 2}[ctx.get("edema", 0)]
    bd["ОКТ: локальные зоны истончения"] = {0: 0, 1: 2}[ctx.get("pachy_uneven", 0)]
    bd["ВГД"] = {"normal": 0, "high": 1, "low": 1}[ctx.get("iog", "normal")]

    min_cat = _cat_min_thickness(int(ctx.get("min_thickness_um", 400)))
    bd["ОКТ: минимальная толщина"] = {">=400": 0, "300-399": 2, "200-299": 4, "<200": 6}[min_cat]
    mean_cat = _cat_mean_thickness(int(ctx.get("mean_thickness_um", 600)))
    bd["ОКТ: средняя толщина"] = {">=600": 0, "520-599": 1, "450-519": 2, "<450": 3}[mean_cat]

    bd["Вовлечение лимба"] = {0: 0, 1: 2}[ctx.get("limbal", 0)]

    bd["Конфокальная: цисты"] = {0: 0, 1: 4, 2: 8, 3: 12}[ctx.get("cysts", 0)]
    bd["Конфокальная: трофозоиты"] = {0: 0, 1: 2}[ctx.get("troph", 0)]
    bd["Конфокальная: глубина цист/трофозоитов"] = {0: 0, 1: 2, 2: 4, 3: 6}[ctx.get("amoeba_depth", 0)]
    bd["Конфокальная: признаки кератоневрита"] = {0: 0, 1: 4}[ctx.get("rk_conf", 0)]

    bd["Длительность до специфической терапии"] = {0: 0, 1: 2, 2: 4}[ctx.get("delay_therapy", 0)]
    bd["Скорость прогрессирования"] = {0: 0, 1: 2, 2: 4}[ctx.get("progress_speed", 0)]
    bd["Прогноз интенсивности помутнения"] = {0: 0, 1: 1, 2: 2}[ctx.get("opacity", 0)]

    score = int(sum(bd.values()))
    critical = (min_cat == "<200") or (ctx.get("depth_cat") == "descemetocele")
    return ScoreResult(score=score, breakdown=bd, critical=critical)


def severity_from_score(score: int, scale: str, critical: bool = False) -> str:
    if critical:
        return "Крайне тяжёлая"
    if scale == "FUSS":
        if score <= 16:
            return "Лёгкая"
        if score <= 26:
            return "Средняя"
        if score <= 36:
            return "Тяжёлая"
        return "Крайне тяжёлая"
    # AUSS
    if score <= 18:
        return "Лёгкая"
    if score <= 30:
        return "Средняя"
    if score <= 42:
        return "Тяжёлая"
    return "Крайне тяжёлая"


def choose_debridement(ctx: Dict[str, Any]) -> str:
    min_um = int(ctx.get("min_thickness_um", 0))
    mean_um = int(ctx.get("mean_thickness_um", 0))
    pachy_uneven = int(ctx.get("pachy_uneven", 0))
    localization = ctx.get("localization", "peripheral")
    total_leucoma = int(ctx.get("total_leucoma", 0))
    edema = int(ctx.get("edema", 0))

    thickness_ok = (min_um >= 400) and (mean_um >= 600) and (pachy_uneven == 0)
    femto_prefer = thickness_ok and (localization in ["paracentral", "central"] or total_leucoma == 1 or edema == 2)
    if femto_prefer:
        return "УФ-кросслинкинг с формированием и удалением роговичного лоскута с использованием фемтосекундного лазера"
    return "УФ-кросслинкинг со скарификацией язвенного инфильтрата под ОКТ-контролем"


def followup_timing(severity: str) -> str:
    if severity == "Лёгкая":
        return "Рекомендуется повторная оценка по шкале через 5–7 суток или раньше при ухудшении."
    if severity == "Средняя":
        return "Рекомендуется повторная оценка по шкале через 48–72 часа."
    if severity == "Тяжёлая":
        return "Рекомендуется повторная оценка по шкале через 24–48 часов."
    return "Повторная оценка по шкале в ближайшие 24 часа/по клиническим показаниям."


def recommend_treatment(scale: str, severity: str, score: int, ctx: Dict[str, Any], critical: bool = False) -> str:
    # Без "самодеятельности": формулировки как в согласованной логике
    if critical or severity == "Крайне тяжёлая":
        return (
            "Экстренная терапевтическая кератопластика.\n"
            "После стойкой стабилизации и элиминации возбудителя возможно рассмотрение отсроченной оптической кератопластики."
        )

    if severity == "Лёгкая":
        base = "Медикаментозная терапия по этиотропной схеме с динамическим контролем."
        return base + "\n" + followup_timing(severity)

    # Средняя/тяжёлая — модифицированный УФ-кросслинкинг
    base = choose_debridement(ctx)
    if severity == "Средняя":
        return f"Модифицированный {base}.\n{followup_timing(severity)}"
    # Тяжёлая
    return (
        f"Модифицированный {base}.\n"
        "При признаках прогрессирующего истончения — переход к экстренной терапевтической кератопластике.\n"
        f"{followup_timing(severity)}"
    )


def format_report_docx_web(scale: str, score: int, severity: str, recommendation: str, breakdown: Dict[str, int] | None = None):
    doc = Document()
    doc.add_heading("Протокол расчёта AUSS/FUSS", level=1)
    doc.add_paragraph(f"Шкала: {scale}")
    doc.add_paragraph(f"Сумма баллов: {score}")
    doc.add_paragraph(f"Степень тяжести: {severity}")
    doc.add_heading("Рекомендации", level=2)
    doc.add_paragraph(recommendation)
    if breakdown:
        doc.add_heading("Разложение баллов", level=2)
        for k, v in breakdown.items():
            doc.add_paragraph(f"{k}: {v}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"{scale}_{date.today().isoformat()}.docx"
    return bio.getvalue(), filename


def format_report_docx_local(patient_name: str, patient_id: str, scale: str, score: int, severity: str, recommendation: str, breakdown: Dict[str, int] | None = None):
    doc = Document()
    doc.add_heading("Протокол расчёта AUSS/FUSS", level=1)
    doc.add_paragraph(f"Данные пациента: {patient_name}".strip())
    if patient_id:
        doc.add_paragraph(f"ID/№карты: {patient_id}".strip())
    doc.add_paragraph(f"Шкала: {scale}")
    doc.add_paragraph(f"Сумма баллов: {score}")
    doc.add_paragraph(f"Степень тяжести: {severity}")
    doc.add_heading("Рекомендации", level=2)
    doc.add_paragraph(recommendation)
    if breakdown:
        doc.add_heading("Разложение баллов", level=2)
        for k, v in breakdown.items():
            doc.add_paragraph(f"{k}: {v}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    safe = (patient_id or patient_name or "patient").replace(" ", "_")
    filename = f"{scale}_{safe}_{date.today().isoformat()}.docx"
    return bio.getvalue(), filename
